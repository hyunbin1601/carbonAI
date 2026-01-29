# RAG (Retrieval-Augmented Generation) 도구
# Chroma DB를 사용한 벡터 검색 및 문서 검색 기능

import os
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any
from react_agent.cache_manager import get_cache_manager

try:
    from langchain_community.vectorstores import Chroma
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document
    from rank_bm25 import BM25Okapi
    import numpy as np
    RAG_AVAILABLE = True
except ImportError:
    RAG_AVAILABLE = False

logger = logging.getLogger(__name__)


class RAGTool:
    """RAG 검색 도구 클래스"""
    
    def __init__(
        self,
        knowledge_base_path: Optional[str] = None,
        chroma_db_path: Optional[str] = None
    ):
        """
        RAG 도구 초기화

        Args:
            knowledge_base_path: 지식베이스 문서 경로
            chroma_db_path: Chroma DB 저장 경로
        """
        if not RAG_AVAILABLE:
            logger.warning("RAG 라이브러리가 설치되지 않았습니다.")
            self.available = False
            return

        self.available = True
        self._kb_last_modified: Optional[float] = None  # 지식베이스 마지막 수정 시간
        
        # 경로 설정
        if knowledge_base_path is None:
            knowledge_base_path = os.getenv(
                "KNOWLEDGE_BASE_PATH",
                str(Path(__file__).parent.parent.parent / "knowledge_base")
            )
        if chroma_db_path is None:
            chroma_db_path = os.getenv(
                "CHROMA_DB_PATH",
                str(Path(__file__).parent.parent.parent / "chroma_db")
            )
        
        self.knowledge_base_path = Path(knowledge_base_path)
        self.chroma_db_path = Path(chroma_db_path)
        
        # 임베딩 모델 초기화
        try:
            # HF_TOKEN 설정 (있으면 사용, 없으면 무시)
            hf_token = os.environ.get("HF_TOKEN")
            if hf_token:
                os.environ["HUGGINGFACE_HUB_TOKEN"] = hf_token

            # Safetensors 자동 변환 비활성화 (타임아웃 방지)
            os.environ["TRANSFORMERS_OFFLINE"] = "0"  # 온라인 유지
            os.environ["HF_HUB_DISABLE_TELEMETRY"] = "1"  # 텔레메트리 비활성화

            self.embeddings = HuggingFaceEmbeddings(
                model_name="dragonkue/BGE-m3-ko",
                model_kwargs={
                    'device': 'cpu',
                    'trust_remote_code': False  # 보안 강화
                },
                encode_kwargs={'normalize_embeddings': True}  # 벡터 정규화 활성화
            )
            logger.info("한국어 임베딩 모델 로드 완료: BGE-m3-ko (1024-dim, 정규화 활성화)")
        except Exception as e:
            logger.warning(f"한국어 임베딩 모델 로드 실패, 기본 모델 사용: {e}")
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                encode_kwargs={'normalize_embeddings': True}  # 벡터 정규화 활성화
            )
        
        # 텍스트 분할기 (의미적 청킹 전략)
        # 한국어 문서에 최적화된 separator 우선순위:
        # 1. 빈 줄 3개 (섹션 구분)
        # 2. 빈 줄 2개 (문단 구분)
        # 3. 빈 줄 1개 (소문단 구분)
        # 4. 한국어/영어 마침표 (문장 단위)
        # 5. 한국어 쉼표, 공백 (어절 단위)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,  # 더 작은 청크로 집중된 컨텍스트 제공
            chunk_overlap=150,  # 20% 오버랩으로 문맥 연속성 유지
            length_function=len,
            separators=[
                "\n\n\n",    # 섹션 구분 (최우선)
                "\n\n",      # 문단 구분
                "\n",        # 줄바꿈
                ". ",        # 영어 마침표
                "。",        # 한국어/중국어 마침표
                "! ",        # 느낌표
                "? ",        # 물음표
                ".",         # 마침표만
                ", ",        # 쉼표
                "，",        # 한국어/중국어 쉼표
                " ",         # 공백 (최후)
                ""           # 문자 단위 (최후의 수단)
            ],
            keep_separator=True  # separator를 유지하여 문맥 보존
        )
        
        # 벡터 스토어 (지연 로딩)
        self._vectorstore: Optional[Chroma] = None

        # BM25 인덱스 (지연 로딩)
        self._bm25_index: Optional['BM25Okapi'] = None
        self._bm25_documents: List[Document] = []  # BM25용 문서 리스트

        # 지식베이스 변경 감지 초기화
        self._update_kb_modified_time()

        logger.info(f"RAG 도구 초기화 완료: {knowledge_base_path}")

    def _get_kb_modified_time(self) -> Optional[float]:
        """지식베이스 디렉토리의 최신 수정 시간 반환"""
        if not self.knowledge_base_path.exists():
            return None

        try:
            latest_time = 0.0
            for ext in ['.txt', '.md', '.pdf', '.docx']:
                for file_path in self.knowledge_base_path.rglob(f"*{ext}"):
                    mtime = file_path.stat().st_mtime
                    if mtime > latest_time:
                        latest_time = mtime
            return latest_time if latest_time > 0 else None
        except Exception as e:
            logger.error(f"지식베이스 수정 시간 확인 실패: {e}")
            return None

    def _update_kb_modified_time(self):
        """지식베이스 수정 시간 업데이트"""
        self._kb_last_modified = self._get_kb_modified_time()
        if self._kb_last_modified:
            logger.info(f"지식베이스 마지막 수정 시간: {self._kb_last_modified}")

    def _check_kb_changed(self) -> bool:
        """지식베이스가 변경되었는지 확인"""
        current_time = self._get_kb_modified_time()
        if current_time is None or self._kb_last_modified is None:
            return False

        changed = current_time > self._kb_last_modified
        if changed:
            logger.info("지식베이스 변경 감지! 캐시를 클리어합니다.")
            # 캐시 클리어
            from react_agent.cache_manager import get_cache_manager
            cache_manager = get_cache_manager()
            cache_manager.clear(prefix="rag")
            cache_manager.clear(prefix="llm")
            # 수정 시간 업데이트
            self._kb_last_modified = current_time

        return changed

    def _extract_section_title(self, content: str, chunk_start_idx: int) -> str:
        """청크가 속한 섹션 제목 추출 (마크다운 헤더 기반)"""
        # 청크 시작 위치 이전의 텍스트에서 가장 가까운 헤더 찾기
        text_before = content[:chunk_start_idx]
        lines = text_before.split('\n')

        # 역순으로 순회하며 마크다운 헤더 찾기
        for line in reversed(lines):
            line = line.strip()
            if line.startswith('#'):
                # 마크다운 헤더 제거하고 제목만 반환
                return line.lstrip('#').strip()

        return ""  # 헤더 없음

    def _extract_keywords_from_text(self, text: str, max_keywords: int = 5) -> List[str]:
        """텍스트에서 핵심 키워드 추출 (간단한 휴리스틱 기반)"""
        # 불용어 제거 및 명사 위주 추출
        stopwords = {'은', '는', '이', '가', '을', '를', '의', '에', '로', '와', '과', '도', '만',
                     '하다', '있다', '되다', '않다', '같다', '위해', '대한', '통해', '따라'}

        # 공백으로 분리하고 불용어 제거
        words = text.split()
        keywords = []

        for word in words:
            # 길이 2 이상, 불용어 아닌 경우
            if len(word) >= 2 and word not in stopwords:
                # 특수문자 제거
                clean_word = ''.join(c for c in word if c.isalnum() or c in ['_', '-'])
                if clean_word and clean_word not in keywords:
                    keywords.append(clean_word)
                    if len(keywords) >= max_keywords:
                        break

        return keywords

    def _load_documents(self) -> List[Document]:
        """지식베이스에서 문서 로드 및 청킹 (메타데이터 강화)"""
        documents = []

        if not self.knowledge_base_path.exists():
            logger.warning(f"지식베이스 경로가 존재하지 않습니다: {self.knowledge_base_path}")
            return documents
        
        # 문서 파싱 함수들
        def parse_text_file(file_path: Path) -> str:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception as e:
                logger.error(f"텍스트 파일 읽기 실패 ({file_path}): {e}")
                return ""
        
        def parse_pdf(file_path: Path) -> str:
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except ImportError:
                logger.warning("pypdf가 설치되지 않았습니다. PDF 파일은 건너뜁니다.")
                return ""
            except Exception as e:
                logger.error(f"PDF 파싱 실패 ({file_path}): {e}")
                return ""
        
        def parse_docx(file_path: Path) -> str:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except ImportError:
                logger.warning("python-docx가 설치되지 않았습니다. DOCX 파일은 건너뜁니다.")
                return ""
            except Exception as e:
                logger.error(f"DOCX 파싱 실패 ({file_path}): {e}")
                return ""
        
        # 지원하는 파일 확장자 및 파서 매핑
        parsers = {
            '.txt': parse_text_file,
            '.md': parse_text_file,
            '.pdf': parse_pdf,
            '.docx': parse_docx,
        }
        
        # 모든 문서 파일 찾기
        for ext, parser_func in parsers.items():
            for file_path in self.knowledge_base_path.rglob(f"*{ext}"):
                try:
                    logger.info(f"문서 로드 중: {file_path.name}")
                    
                    # 문서 파싱
                    content = parser_func(file_path)
                    if not content.strip():
                        logger.warning(f"빈 문서: {file_path.name}")
                        continue
                    
                    # 텍스트 분할
                    chunks = self.text_splitter.split_text(content)

                    # Document 객체 생성 (강화된 메타데이터)
                    for i, chunk in enumerate(chunks):
                        # 청크 위치 정보
                        if len(chunks) == 1:
                            position = "full"
                        elif i == 0:
                            position = "beginning"
                        elif i == len(chunks) - 1:
                            position = "end"
                        else:
                            position = "middle"

                        # 섹션 제목 추출 (청크 내 첫 헤더 또는 첫 줄)
                        section_title = ""
                        chunk_lines = chunk.split('\n')
                        for line in chunk_lines[:3]:  # 처음 3줄 확인
                            line = line.strip()
                            if line.startswith('#'):
                                section_title = line.lstrip('#').strip()
                                break

                        # 키워드 추출 (청크 내용에서)
                        keywords = self._extract_keywords_from_text(chunk, max_keywords=5)

                        doc = Document(
                            page_content=chunk,
                            metadata={
                                'source': str(file_path),
                                'filename': file_path.name,
                                'extension': ext,
                                'chunk_index': i,
                                'total_chunks': len(chunks),
                                'position': position,  # beginning/middle/end/full
                                'chunk_length': len(chunk),
                                'section_title': section_title,  # 섹션 제목
                                'keywords': ', '.join(keywords),  # 키워드 (쉼표로 구분)
                            }
                        )
                        documents.append(doc)
                    
                    logger.info(f"✓ 문서 로드 완료: {file_path.name} ({len(chunks)}개 청크)")
                    
                except Exception as e:
                    logger.error(f"문서 로드 실패 ({file_path}): {e}")
                    continue
        
        logger.info(f"총 {len(documents)}개 문서 청크 로드 완료")
        return documents
    
    def _build_vectorstore_if_needed(self) -> bool:
        """벡터 DB가 없으면 자동으로 구축"""
        # 이미 벡터 스토어가 있으면 건너뛰기
        if self._vectorstore is not None:
            return True
        
        # 벡터 DB가 이미 존재하면 로드만
        if self.chroma_db_path.exists() and any(self.chroma_db_path.iterdir()):
            return False
        
        # 지식베이스에 문서가 있는지 확인
        if not self.knowledge_base_path.exists():
            logger.warning(f"지식베이스 경로가 존재하지 않습니다: {self.knowledge_base_path}")
            return False
        
        # 문서 찾기
        has_documents = False
        for ext in ['.txt', '.md', '.pdf', '.docx']:
            if any(self.knowledge_base_path.rglob(f"*{ext}")):
                has_documents = True
                break
        
        if not has_documents:
            logger.warning("지식베이스에 문서가 없습니다. 벡터 DB를 구축할 수 없습니다.")
            return False
        
        # 벡터 DB 자동 구축
        logger.info("벡터 DB가 없습니다. 자동으로 구축을 시작합니다...")
        try:
            documents = self._load_documents()
            
            if not documents:
                logger.warning("로드할 문서가 없습니다.")
                return False
            
            # Chroma DB 생성 (HNSW 최적화 + cosine distance 사용)
            logger.info(f"벡터 DB 구축 중... ({len(documents)}개 문서 청크)")
            self._vectorstore = Chroma.from_documents(
                documents=documents,
                embedding=self.embeddings,
                persist_directory=str(self.chroma_db_path),
                collection_metadata={
                    "hnsw:space": "cosine",  # Cosine distance 명시
                    "hnsw:M": 16,  # 기본값(16)보다 낮춤 - 속도 우선 (메모리/속도 트레이드오프)
                    "hnsw:ef_construction": 100,  # 인덱스 구축 시 탐색 깊이 (품질 유지)
                    "hnsw:ef": 50  # 검색 시 탐색 깊이 (기본값보다 낮춤 - 속도 우선)
                }
            )

            # 거리 메트릭 검증
            try:
                actual_metric = self._vectorstore._collection.metadata.get('hnsw:space', 'unknown')
                logger.info(f"✓ 벡터 DB 구축 완료: {len(documents)}개 문서")
                logger.info(f"  - Distance metric: {actual_metric}")
                logger.info(f"  - Embeddings normalized: True")
                if actual_metric != 'cosine':
                    logger.warning(f"예상 메트릭(cosine)과 실제({actual_metric})가 다릅니다!")
            except Exception as e:
                logger.warning(f"거리 메트릭 검증 실패: {e}")

            return True
            
        except Exception as e:
            logger.error(f"벡터 DB 구축 실패: {e}")
            return False
    
    @property
    def vectorstore(self) -> Optional[Chroma]:
        """벡터 스토어 지연 로딩 및 자동 구축"""
        if self._vectorstore is None and self.available:
            try:
                # 벡터 DB가 없으면 자동 구축 시도
                self._build_vectorstore_if_needed()

                # 기존 벡터 DB 로드 또는 새로 구축된 DB 사용
                if self.chroma_db_path.exists() and any(self.chroma_db_path.iterdir()):
                    if self._vectorstore is None:
                        self._vectorstore = Chroma(
                            persist_directory=str(self.chroma_db_path),
                            embedding_function=self.embeddings
                        )

                    # 진단: ChromaDB distance 함수 확인
                    try:
                        collection = self._vectorstore._collection
                        metadata = collection.metadata
                        distance_function = metadata.get('hnsw:space', 'unknown')
                        logger.info(f"ChromaDB distance 함수: {distance_function}")
                    except Exception as e:
                        logger.warning(f" Distance 함수 확인 실패: {e}")

                    logger.info("벡터 DB 로드 완료")
                else:
                    logger.warning("벡터 DB가 아직 구축되지 않았습니다.")
            except Exception as e:
                logger.error(f"벡터 스토어 로드 실패: {e}")

        return self._vectorstore

    def _normalize_query(self, query: str) -> str:
        """검색 쿼리 정규화 (최소한의 정규화만 수행)"""
        import re

        # 1. 연속된 공백을 하나로 통일
        normalized = re.sub(r'\s+', ' ', query)

        # 2. 앞뒤 공백 제거
        normalized = normalized.strip()

        # 참고: 띄어쓰기 패턴 정규화는 제거
        # 이유: 임베딩 모델이 이미 띄어쓰기를 잘 처리하며,
        #       원본 문서의 표기법을 존중하는 것이 더 나은 결과를 제공

        return normalized

    def _tokenize(self, text: str) -> List[str]:
        """텍스트를 토큰으로 분할 (한국어/영어 지원)"""
        # 간단한 공백 기반 토크나이징
        # 더 정교한 토크나이징이 필요하면 konlpy 등 사용 가능
        import re
        # 특수문자 제거 및 소문자 변환
        text = text.lower()
        # 단어 단위로 분리 (한글, 영문, 숫자만)
        tokens = re.findall(r'[가-힣a-z0-9]+', text)
        return tokens

    def _build_bm25_index(self) -> bool:
        """BM25 인덱스 빌드"""
        if not self.available:
            return False

        try:
            # 벡터 DB가 있으면 거기서 문서 가져오기
            if self.vectorstore is not None:
                # Chroma DB에서 모든 문서 가져오기
                collection = self.vectorstore._collection
                results = collection.get(include=['documents', 'metadatas'])

                if not results or not results['documents']:
                    logger.warning("BM25 인덱스 구축 실패: 문서가 없습니다.")
                    return False

                # Document 객체로 변환
                self._bm25_documents = []
                for i, (doc_text, metadata) in enumerate(zip(results['documents'], results['metadatas'])):
                    doc = Document(page_content=doc_text, metadata=metadata or {})
                    self._bm25_documents.append(doc)

            else:
                # 벡터 DB가 없으면 직접 로드
                self._bm25_documents = self._load_documents()

            if not self._bm25_documents:
                logger.warning("BM25 인덱스 구축 실패: 문서가 없습니다.")
                return False

            # 각 문서를 토큰화
            tokenized_corpus = [
                self._tokenize(doc.page_content)
                for doc in self._bm25_documents
            ]

            # BM25 인덱스 생성
            self._bm25_index = BM25Okapi(tokenized_corpus)
            logger.info(f"✓ BM25 인덱스 구축 완료: {len(self._bm25_documents)}개 문서")
            return True

        except Exception as e:
            logger.error(f"BM25 인덱스 구축 실패: {e}")
            return False

    @property
    def bm25_index(self) -> Optional['BM25Okapi']:
        """BM25 인덱스 지연 로딩"""
        if self._bm25_index is None and self.available:
            self._build_bm25_index()
        return self._bm25_index

    def search_documents(self, query: str, k: int = 3, similarity_threshold: float = 0.5) -> List[Dict[str, Any]]:
        """
        문서 검색 (키워드 추출 + 코사인 유사도 기반)

        Args:
            query: 검색 쿼리
            k: 반환할 문서 수
            similarity_threshold: 코사인 유사도 임계값 (기본값: 0.5, 한국어 모델에 적합)

        Returns:
            관련 문서 리스트 (딕셔너리 형태)
        """
        if not self.available or self.vectorstore is None:
            return []

        # 지식베이스 변경 확인 (변경되면 자동으로 캐시 클리어됨)
        self._check_kb_changed()

        # 캐시 확인 (쿼리 + 파라미터 조합으로 캐싱)
        cache_manager = get_cache_manager()
        cache_content = f"{query}|k={k}|threshold={similarity_threshold}"
        cached_result = cache_manager.get("rag", cache_content)
        if cached_result is not None:
            return cached_result

        try:
            # 쿼리 정규화 (일관성 향상)
            normalized_query = self._normalize_query(query)
            if normalized_query != query:
                logger.info(f"[검색] 쿼리 정규화: '{query}' → '{normalized_query}'")
            else:
                logger.info(f"[검색] 쿼리: '{query}'")

            # 지식베이스 문서 수 확인
            try:
                total_docs = self.vectorstore._collection.count()
                logger.info(f"지식베이스: 총 {total_docs}개 문서 청크")
            except Exception as e:
                logger.warning(f"지식베이스 문서 수 확인 실패: {e}")

            # 정규화된 쿼리로 검색 (일관성 + 품질)
            docs_with_scores = self.vectorstore.similarity_search_with_score(normalized_query, k=k * 3)
            logger.info(f"벡터 검색: {len(docs_with_scores)}개 결과")

            # 상위 5개 결과의 실제 점수 출력 (임계값 필터링 전)
            logger.debug(f"상위 {min(5, len(docs_with_scores))}개 문서 (필터링 전):")
            for idx, (doc, distance) in enumerate(docs_with_scores[:5]):
                similarity = 1.0 - distance if distance <= 2.0 else max(0.0, 1.0 - (distance / 2.0))
                filename = doc.metadata.get('filename', 'unknown')
                preview = doc.page_content[:50].replace('\n', ' ')
                logger.debug(f"  #{idx+1}: {filename} (거리: {distance:.4f}, 유사도: {similarity:.4f}) - {preview}...")

            filtered_docs = []
            seen_keys = set()  # 중복 제거용 (source + chunk_index 조합)
            rejected_count = 0

            for idx, (doc, distance) in enumerate(docs_with_scores):
                # Chroma DB의 distance 처리
                # Chroma는 L2 거리 또는 코사인 거리를 사용할 수 있음
                # 정규화된 벡터의 경우 코사인 거리: 0 ~ 2 (값이 작을수록 유사)
                # 코사인 유사도 = 1 - 코사인 거리

                # distance가 음수이거나 2보다 크면 이상한 값이므로 처리
                if distance < 0:
                    similarity = 1.0  # 완전 일치로 간주
                elif distance > 2.0:
                    similarity = max(0.0, 1.0 - (distance / 2.0))  # 정규화
                else:
                    similarity = 1.0 - distance

                # 유사도 임계값 확인
                if similarity < similarity_threshold:
                    rejected_count += 1
                    continue
                
                # 중복 제거: 같은 source와 chunk_index 조합은 제외
                source = doc.metadata.get('source', 'unknown')
                chunk_index = doc.metadata.get('chunk_index', 0)
                doc_key = (source, chunk_index)
                
                if doc_key in seen_keys:
                    continue
                
                seen_keys.add(doc_key)
                
                filtered_docs.append({
                    'content': doc.page_content,
                    'source': source,
                    'filename': doc.metadata.get('filename', 'unknown'),
                    'chunk_index': chunk_index,
                    'similarity': float(similarity)  # Convert to native Python float
                })

                # 최대 k개까지만 수집 (유사도가 높은 순서대로)
                if len(filtered_docs) >= k:
                    break

            if not filtered_docs:
                logger.warning(
                    f"임계값 {similarity_threshold} 미만: "
                    f"{len(docs_with_scores)}개 결과 모두 제외됨"
                )
                # 빈 결과도 캐싱 (불필요한 재검색 방지, TTL은 짧게)
                cache_manager.set("rag", cache_content, [], ttl=3600)  # 1시간
                return []

            logger.info(f"필터링 완료: {len(filtered_docs)}개 선택, {rejected_count}개 제외")

            # 임계값을 넘긴 문서들의 유사도 로깅
            logger.info(f"최종 결과: {len(filtered_docs)}개")
            for idx, doc in enumerate(filtered_docs[:5]):  # 상위 5개만 출력
                logger.info(f"  #{idx+1}: {doc['filename']} (유사도: {doc['similarity']:.3f})")

            # 검색 결과 캐싱 (24시간)
            cache_manager.set("rag", cache_content, filtered_docs)

            return filtered_docs
            
        except Exception as e:
            logger.error(f"[검색] 검색 실패: {e}", exc_info=True)
            return []

    def search_documents_hybrid(
        self,
        query: str,
        k: int = 3,
        alpha: float = 0.5,
        similarity_threshold: float = 0.5
    ) -> List[Dict[str, Any]]:
        """
        하이브리드 검색 (BM25 + 벡터 검색)

        Args:
            query: 검색 쿼리
            k: 반환할 문서 수
            alpha: 벡터 검색 가중치 (0~1, 기본값 0.5)
                   - 1.0: 100% 벡터 검색
                   - 0.0: 100% BM25 검색
                   - 0.5: 50% 벡터 + 50% BM25
            similarity_threshold: 최소 유사도 임계값

        Returns:
            관련 문서 리스트 (하이브리드 점수로 정렬)
        """
        if not self.available:
            return []

        # 지식베이스 변경 확인
        self._check_kb_changed()

        # 캐시 확인
        cache_manager = get_cache_manager()
        cache_content = f"hybrid|{query}|k={k}|alpha={alpha}|threshold={similarity_threshold}"
        cached_result = cache_manager.get("rag", cache_content)
        if cached_result is not None:
            return cached_result

        try:
            # 쿼리 정규화 (일관성 향상)
            normalized_query = self._normalize_query(query)
            if normalized_query != query:
                logger.info(f"[하이브리드] 쿼리 정규화: '{query}' → '{normalized_query}'")
            else:
                logger.info(f"[하이브리드] 쿼리: '{query}'")

            # 지식베이스 문서 수 확인
            try:
                total_docs = self.vectorstore._collection.count()
                logger.info(f"지식베이스: 총 {total_docs}개 문서 청크")
            except Exception as e:
                logger.warning(f"지식베이스 문서 수 확인 실패: {e}")

            # 1. 벡터 검색
            vector_results = {}
            if self.vectorstore is not None:
                vector_docs = self.vectorstore.similarity_search_with_score(normalized_query, k=k * 3)
                print(f"[VECTOR] {len(vector_docs)} results")

                # 진단: 실제 distance 값 확인
                if vector_docs:
                    print(f"[VECTOR] Top 3 (distance → similarity):")
                    for idx, (doc, distance) in enumerate(vector_docs[:3]):
                        filename = doc.metadata.get('filename', 'unknown')
                        similarity = 1.0 - distance
                        status = "HIGH" if similarity >= 0.7 else "MED" if similarity >= 0.5 else "LOW"
                        print(f"  - {filename[:60]}: dist={distance:.4f} → sim={similarity:.4f} [{status}]")

                for doc, distance in vector_docs:
                    doc_key = (doc.metadata.get('source', ''), doc.metadata.get('chunk_index', 0))
                    # 거리를 유사도로 변환 (0~1 범위)
                    similarity = max(0.0, 1.0 - min(distance, 2.0))
                    vector_results[doc_key] = {
                        'doc': doc,
                        'score': similarity
                    }

            # 2. BM25 검색
            bm25_results = {}
            if self.bm25_index is not None and len(self._bm25_documents) > 0:
                # 쿼리 토크나이징 (정규화된 쿼리 사용)
                tokenized_query = self._tokenize(normalized_query)
                # BM25 점수 계산
                bm25_scores = self.bm25_index.get_scores(tokenized_query)

                # 점수 정규화 (0~1 범위로)
                max_score = max(bm25_scores) if len(bm25_scores) > 0 and max(bm25_scores) > 0 else 1.0
                normalized_scores = [score / max_score for score in bm25_scores]

                # 상위 k*3개만 선택
                top_indices = np.argsort(normalized_scores)[::-1][:k * 3]
                print(f"[BM25] {len(top_indices)} results")

                for idx in top_indices:
                    doc = self._bm25_documents[idx]
                    doc_key = (doc.metadata.get('source', ''), doc.metadata.get('chunk_index', 0))
                    bm25_results[doc_key] = {
                        'doc': doc,
                        'score': normalized_scores[idx]
                    }

            # 3. 점수 결합 (RRF - Reciprocal Rank Fusion 또는 가중 평균)
            combined_results = {}
            all_doc_keys = set(vector_results.keys()) | set(bm25_results.keys())

            # RRF 방식: 순위 기반 융합 (더 robust)
            # 참고: alpha < 0 이면 RRF 사용, alpha >= 0 이면 가중 평균 사용
            use_rrf = alpha < 0  # 음수 alpha는 RRF 활성화 신호

            if use_rrf:
                # RRF 방식: 순위로 정렬, 점수는 vector+bm25 평균 사용
                k_rrf = 60  # RRF 상수 (일반적으로 60이 최적)

                # Vector 결과 순위 생성
                vector_sorted = sorted(
                    vector_results.items(),
                    key=lambda x: x[1]['score'],
                    reverse=True
                )
                vector_ranks = {doc_key: rank for rank, (doc_key, _) in enumerate(vector_sorted)}

                # BM25 결과 순위 생성
                bm25_sorted = sorted(
                    bm25_results.items(),
                    key=lambda x: x[1]['score'],
                    reverse=True
                )
                bm25_ranks = {doc_key: rank for rank, (doc_key, _) in enumerate(bm25_sorted)}

                # RRF로 순위 계산 및 결과 저장 (Numpy vectorization으로 최적화)
                all_doc_keys_list = list(all_doc_keys)

                # Vectorized rank extraction
                vector_ranks_array = np.array([
                    vector_ranks.get(doc_key, len(vector_results))
                    for doc_key in all_doc_keys_list
                ])
                bm25_ranks_array = np.array([
                    bm25_ranks.get(doc_key, len(bm25_results))
                    for doc_key in all_doc_keys_list
                ])

                # Vectorized RRF calculation: 1/(k + rank + 1)
                rrf_rank_scores = (1.0 / (k_rrf + vector_ranks_array + 1) +
                                   1.0 / (k_rrf + bm25_ranks_array + 1))

                # Vectorized score extraction
                vector_scores = np.array([
                    vector_results.get(doc_key, {}).get('score', 0.0)
                    for doc_key in all_doc_keys_list
                ])
                bm25_scores = np.array([
                    bm25_results.get(doc_key, {}).get('score', 0.0)
                    for doc_key in all_doc_keys_list
                ])

                # 최종 점수: vector와 bm25의 평균 (절대적 품질 유지)
                hybrid_scores = (vector_scores + bm25_scores) / 2.0

                # 결과 저장
                for idx, doc_key in enumerate(all_doc_keys_list):
                    doc = vector_results.get(doc_key, bm25_results.get(doc_key, {})).get('doc')

                    if doc is not None:
                        combined_results[doc_key] = {
                            'doc': doc,
                            'rrf_rank_score': float(rrf_rank_scores[idx]),  # 정렬용
                            'hybrid_score': float(hybrid_scores[idx]),  # 임계값 필터링용
                            'vector_score': float(vector_scores[idx]),
                            'bm25_score': float(bm25_scores[idx])
                        }
            else:
                # 가중 평균 방식 (기존)
                for doc_key in all_doc_keys:
                    vector_score = vector_results.get(doc_key, {}).get('score', 0.0)
                    bm25_score = bm25_results.get(doc_key, {}).get('score', 0.0)

                    # 하이브리드 점수 계산
                    hybrid_score = alpha * vector_score + (1.0 - alpha) * bm25_score

                    # 문서 객체 가져오기 (벡터 우선, 없으면 BM25)
                    doc = vector_results.get(doc_key, bm25_results.get(doc_key, {})).get('doc')

                    if doc is not None:
                        combined_results[doc_key] = {
                            'doc': doc,
                            'hybrid_score': hybrid_score,
                            'vector_score': vector_score,
                            'bm25_score': bm25_score
                        }

            # 4. 하이브리드 점수로 정렬
            # RRF 사용 시: rrf_rank_score로 정렬, 가중평균 사용 시: hybrid_score로 정렬
            if use_rrf:
                sorted_results = sorted(
                    combined_results.items(),
                    key=lambda x: x[1].get('rrf_rank_score', x[1]['hybrid_score']),
                    reverse=True
                )
            else:
                sorted_results = sorted(
                    combined_results.items(),
                    key=lambda x: x[1]['hybrid_score'],
                    reverse=True
                )

            # 상위 5개 결과의 실제 점수 출력 (임계값 필터링 전)
            logger.debug(f"상위 {min(5, len(sorted_results))}개 문서 (필터링 전):")
            for idx, (doc_key, result) in enumerate(sorted_results[:5]):
                doc = result['doc']
                filename = doc.metadata.get('filename', 'unknown')
                preview = doc.page_content[:50].replace('\n', ' ')
                logger.debug(
                    f"  #{idx+1}: {filename} "
                    f"(hybrid: {result['hybrid_score']:.3f} = "
                    f"vector: {result['vector_score']:.3f} + bm25: {result['bm25_score']:.3f}) "
                    f"- {preview}..."
                )

            # 5. 필터링 및 결과 생성
            filtered_docs = []
            rejected_count = 0
            for idx, (doc_key, result) in enumerate(sorted_results):
                hybrid_score = result['hybrid_score']

                # 임계값 확인
                if hybrid_score < similarity_threshold:
                    rejected_count += 1
                    continue

                doc = result['doc']
                filtered_docs.append({
                    'content': doc.page_content,
                    'source': doc.metadata.get('source', 'unknown'),
                    'filename': doc.metadata.get('filename', 'unknown'),
                    'chunk_index': doc.metadata.get('chunk_index', 0),
                    'similarity': float(hybrid_score),  # Convert to native Python float
                    'vector_score': float(result['vector_score']),
                    'bm25_score': float(result['bm25_score'])
                })

                # 최대 k개까지
                if len(filtered_docs) >= k:
                    break

            if not filtered_docs:
                logger.warning(
                    f"임계값 {similarity_threshold} 미만: "
                    f"{len(sorted_results)}개 결과 모두 제외됨"
                )
                cache_manager.set("rag", cache_content, [], ttl=3600)
                return []

            logger.info(f"필터링 완료: {len(filtered_docs)}개 선택, {rejected_count}개 제외")

            # 최종 결과 로깅
            logger.info(f"최종 결과 (하이브리드, alpha={alpha}): {len(filtered_docs)}개")
            for idx, doc in enumerate(filtered_docs[:5]):  # 상위 5개만 출력
                logger.info(
                    f"  #{idx+1}: {doc['filename']} "
                    f"(hybrid: {doc['similarity']:.3f} = vector: {doc['vector_score']:.3f} + bm25: {doc['bm25_score']:.3f})"
                )

            # 검색 결과 캐싱
            cache_manager.set("rag", cache_content, filtered_docs)
            return filtered_docs

        except Exception as e:
            logger.error(f"[하이브리드] 검색 실패: {e}", exc_info=True)
            return []


# 전역 RAG 도구 인스턴스
_rag_tool: Optional[RAGTool] = None


def get_rag_tool() -> RAGTool:
    """RAG 도구 싱글톤 인스턴스 반환"""
    global _rag_tool
    if _rag_tool is None:
        _rag_tool = RAGTool()
    return _rag_tool

