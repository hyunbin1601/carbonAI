"""
벡터 DB 구축 스크립트
지식베이스 문서를 로드하여 Chroma DB에 저장합니다.
"""

import os
import logging
from pathlib import Path
from typing import List
from dotenv import load_dotenv

# 환경 변수 로드
load_dotenv()

# 로깅 설정
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

try:
    from langchain_community.vectorstores import Chroma
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document
    RAG_AVAILABLE = True
except ImportError as e:
    logger.error(f"필수 라이브러리가 설치되지 않았습니다: {e}")
    logger.error("다음 명령어로 설치하세요: pip install chromadb sentence-transformers langchain-community langchain-text-splitters")
    RAG_AVAILABLE = False
    Document = None  # 타입 힌트를 위한 더미

# 문서 파싱 라이브러리
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx가 설치되지 않았습니다. DOCX 파일은 건너뜁니다.")

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("pypdf가 설치되지 않았습니다. PDF 파일은 건너뜁니다.")


def parse_text_file(file_path: Path) -> str:
    """텍스트 파일 파싱"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.error(f"텍스트 파일 읽기 실패 ({file_path}): {e}")
        return ""


def parse_markdown(file_path: Path) -> str:
    """마크다운 파일 파싱"""
    return parse_text_file(file_path)


def parse_pdf(file_path: Path) -> str:
    """PDF 파일 파싱"""
    if not PDF_AVAILABLE:
        return ""
    
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"PDF 파싱 실패 ({file_path}): {e}")
        return ""


def parse_docx(file_path: Path) -> str:
    """DOCX 파일 파싱"""
    if not DOCX_AVAILABLE:
        return ""
    
    try:
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"DOCX 파싱 실패 ({file_path}): {e}")
        return ""


def load_documents(knowledge_base_path: Path):
    """
    지식베이스에서 모든 문서를 로드하고 청킹
    
    Returns:
        Document 객체 리스트
    """
    documents = []
    
    if not knowledge_base_path.exists():
        logger.warning(f"지식베이스 경로가 존재하지 않습니다: {knowledge_base_path}")
        return documents
    
    # 지원하는 파일 확장자 및 파서 매핑
    parsers = {
        '.txt': parse_text_file,
        '.md': parse_markdown,
        '.pdf': parse_pdf,
        '.docx': parse_docx,
    }
    
    # 모든 문서 파일 찾기
    for ext, parser_func in parsers.items():
        for file_path in knowledge_base_path.rglob(f"*{ext}"):
            try:
                logger.info(f"문서 로드 중: {file_path.name}")
                
                # 문서 파싱
                content = parser_func(file_path)
                if not content.strip():
                    logger.warning(f"빈 문서: {file_path.name}")
                    continue
                
                # 텍스트 분할 (rag_tool.py와 동일하게 통일)
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=900,  # ✅ 통일: 절차 + 개념 설명 적정 크기
                    chunk_overlap=180,  # ✅ 통일: 20% (문맥 연속성)
                    length_function=len,
                    separators=[
                        "\n\n\n",    # 섹션 구분 (최우선)
                        "\n\n",      # 문단 구분
                        "\n",        # 줄바꿈
                        ". ",        # 영어 마침표
                        "。",        # 한국어 마침표
                        "! ",        # 느낌표
                        "? ",        # 물음표
                        ".",         # 마침표만
                        ", ",        # 쉼표
                        "，",        # 한국어 쉼표
                        " ",         # 공백 (최후)
                        ""           # 문자 단위 (최후의 수단)
                    ],
                    keep_separator=True  # separator 유지로 문맥 보존
                )
                chunks = text_splitter.split_text(content)
                
                # Document 객체 생성
                for i, chunk in enumerate(chunks):
                    doc = Document(
                        page_content=chunk,
                        metadata={
                            'source': str(file_path),
                            'filename': file_path.name,
                            'extension': ext,
                            'chunk_index': i,
                            'total_chunks': len(chunks)
                        }
                    )
                    documents.append(doc)
                
                logger.info(f"✓ 문서 로드 완료: {file_path.name} ({len(chunks)}개 청크)")
                
            except Exception as e:
                logger.error(f"문서 로드 실패 ({file_path}): {e}")
                continue
    
    logger.info(f"총 {len(documents)}개 문서 청크 로드 완료")
    return documents


def build_vectorstore(force_rebuild: bool = False):
    """
    벡터 스토어 구축
    
    Args:
        force_rebuild: 기존 DB를 삭제하고 재구축할지 여부
    """
    if not RAG_AVAILABLE:
        logger.error("RAG 라이브러리가 설치되지 않았습니다.")
        return
    
    # 경로 설정
    base_path = Path(__file__).parent
    knowledge_base_path = Path(os.getenv(
        "KNOWLEDGE_BASE_PATH",
        str(base_path / "knowledge_base")
    ))
    chroma_db_path = Path(os.getenv(
        "CHROMA_DB_PATH",
        str(base_path / "chroma_db")
    ))
    
    logger.info(f"지식베이스 경로: {knowledge_base_path}")
    logger.info(f"Chroma DB 경로: {chroma_db_path}")
    
    # 기존 DB 삭제 (force_rebuild가 True인 경우)
    if force_rebuild and chroma_db_path.exists():
        import shutil
        shutil.rmtree(chroma_db_path)
        logger.info("기존 벡터 DB 삭제 완료")
    
    # 기존 벡터 스토어가 있으면 건너뛰기
    if chroma_db_path.exists() and any(chroma_db_path.iterdir()):
        logger.info("기존 벡터 DB가 이미 존재합니다. 재구축하려면 --force 옵션을 사용하세요.")
        return
    
    # 임베딩 모델 초기화
    try:
        embeddings = HuggingFaceEmbeddings(
            model_name="jhgan/ko-sroberta-multitask",
            model_kwargs={'device': 'cpu'}
        )
        logger.info("한국어 임베딩 모델 로드 완료")
    except Exception as e:
        logger.warning(f"한국어 임베딩 모델 로드 실패, 기본 모델 사용: {e}")
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
    
    # 문서 로드
    documents = load_documents(knowledge_base_path)
    
    if not documents:
        logger.warning("로드할 문서가 없습니다. knowledge_base 폴더에 문서를 추가하세요.")
        logger.info("지원 형식: .txt, .md, .pdf, .docx")
        return
    
    # Chroma DB 생성 (코사인 거리 메트릭 사용)
    logger.info("벡터 DB 구축 중... (시간이 걸릴 수 있습니다)")
    vectorstore = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory=str(chroma_db_path),
        collection_metadata={"hnsw:space": "cosine"}  # 코사인 거리 메트릭 사용
    )
    
    logger.info(f"✓ 벡터 스토어 구축 완료: {len(documents)}개 문서")
    logger.info(f"✓ 저장 위치: {chroma_db_path}")


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="벡터 DB 구축 스크립트")
    parser.add_argument(
        "--force",
        action="store_true",
        help="기존 DB를 삭제하고 재구축"
    )
    
    args = parser.parse_args()
    
    build_vectorstore(force_rebuild=args.force)

